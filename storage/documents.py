"""Extract plain text from imported documents (.pdf, .docx, .txt/.md)."""

from pathlib import Path

SUPPORTED_EXTENSIONS = (".txt", ".md", ".pdf", ".docx")

# File-dialog filter string for Qt.
FILE_FILTER = "Documents (*.pdf *.docx *.txt *.md);;All files (*)"


def _extract_txt(path: Path) -> str:
    # utf-8-sig strips a leading byte-order mark if present (Notepad / PowerShell
    # write one), so the text doesn't start with a stray U+FEFF character.
    return path.read_text(encoding="utf-8-sig", errors="ignore")


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]

    # include table cell text, which holds real content in many resumes
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text(path) -> str:
    """Return the plain-text content of a document, dispatched by extension."""
    path = Path(path)
    ext = path.suffix.lower()

    if ext in (".txt", ".md"):
        text = _extract_txt(path)
    elif ext == ".pdf":
        text = _extract_pdf(path)
    elif ext == ".docx":
        text = _extract_docx(path)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    return text.strip()
